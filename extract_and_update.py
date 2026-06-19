#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
import json
import re
from collections import OrderedDict

def extract_content_from_docx(filename, grade):
    """docx 파일에서 문제와 답을 추출"""
    doc = Document(filename)
    
    questions = []
    answers = {}
    
    # 모든 텍스트 수집
    all_text = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            all_text.append(text)
    
    # 테이블에서 정답 추출
    print(f"\n=== {filename} 테이블 분석 ===")
    for table_idx, table in enumerate(doc.tables):
        print(f"\n테이블 #{table_idx + 1}")
        print(f"행: {len(table.rows)}, 열: {len(table.columns)}")
        
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            print(f"  행 {i}: {cells}")
            
            # 문제 번호와 답 추출 (일반적인 형식: [문제번호] [답])
            if len(cells) >= 2:
                first_cell = cells[0].strip()
                second_cell = cells[1].strip()
                
                # 숫자 추출
                match = re.match(r'^(\d+)', first_cell)
                if match:
                    q_num = int(match.group(1))
                    answer = second_cell
                    answers[q_num] = answer
                    print(f"    → 문제 #{q_num}: {answer}")
    
    print(f"\n추출된 정답 수: {len(answers)}")
    return questions, answers

# 3학년 파일 처리
print("="*60)
print("3학년 광고콘텐츠 재작.docx 분석")
print("="*60)
q3, a3 = extract_content_from_docx('(3학년) 광고콘텐츠 재작.docx', 3)

# 2학년 파일 처리
print("\n" + "="*60)
print("2학년 광고콘텐츠제작.docx 분석")
print("="*60)
q2, a2 = extract_content_from_docx('2학년 광고콘텐츠제작.docx', 2)

# 결과 저장
print("\n\n=== 추출 결과 ===")
print(f"3학년 정답: {a3}")
print(f"2학년 정답: {a2}")

# 정답을 JSON으로 저장
with open('extracted_answers.json', 'w', encoding='utf-8') as f:
    json.dump({
        '3rd_grade': a3,
        '2nd_grade': a2
    }, f, ensure_ascii=False, indent=2)

print("\nextracted_answers.json 파일 생성 완료")
