#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document

# 3학년 문서 읽기
doc = Document('(3학년) 광고콘텐츠 재작.docx')

print("=== 테이블 개수 ===")
print(f"테이블: {len(doc.tables)}개\n")

# 테이블 내용 출력
for table_idx, table in enumerate(doc.tables):
    print(f"\n=== 테이블 {table_idx + 1} ===")
    print(f"행: {len(table.rows)}, 열: {len(table.columns)}")
    
    # 처음 20행만 표시
    for i, row in enumerate(table.rows[:20]):
        cells = [cell.text.strip() for cell in row.cells]
        print(" | ".join(cells))

print("\n\n=== 14번, 46번 정답 찾기 ===\n")

for table_idx, table in enumerate(doc.tables):
    for i, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]
        if len(cells) >= 1:
            if cells[0] == '14' or cells[0] == '46':
                print(f"찾음! 테이블 {table_idx} 행 {i}")
                print(f"내용: {cells}")
