#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.table import Table
import re

# 3학년 문서 읽기
doc = Document('(3학년) 광고콘텐츠 재작.docx')

print("=== DOCX 문서의 테이블 찾기 ===\n")

table_count = 0
for i, element in enumerate(doc.element.body):
    if element.tag.endswith('tbl'):
        table_count += 1
        print(f"\n테이블 #{table_count} 발견 (위치: {i})")
        
        # 테이블을 Table 객체로 변환
        from docx.oxml import parse_xml
        
        # 테이블의 모든 행 출력
        tables = doc.tables
        if table_count <= len(tables):
            tbl = tables[table_count - 1]
            print(f"행 수: {len(tbl.rows)}, 열 수: {len(tbl.columns)}")
            
            # 처음 10행 출력
            for row_idx, row in enumerate(tbl.rows[:10]):
                cells_content = [cell.text.strip() for cell in row.cells]
                print(f"  행 {row_idx}: {cells_content}")

print(f"\n\n총 {len(doc.tables)}개 테이블 발견")

# 모든 테이블 출력
for t_idx, table in enumerate(doc.tables):
    print(f"\n=== 테이블 {t_idx+1} ===")
    print(f"행: {len(table.rows)}, 열: {len(table.columns)}")
    
    for row_idx, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]
        print(f"행 {row_idx}: {cells}")
