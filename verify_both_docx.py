#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
import json

print("="*60)
print("3학년 광고콘텐츠 재작.docx 정답표")
print("="*60)

doc3 = Document('(3학년) 광고콘텐츠 재작.docx')

# 3학년 정답표 추출
for table in doc3.tables:
    print(f"\n테이블 발견 (행: {len(table.rows)}, 열: {len(table.columns)})\n")
    
    for i, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]
        print(f"{cells}")
        
print("\n\n" + "="*60)
print("2학년 광고콘텐츠제작.docx 정답표")
print("="*60)

try:
    doc2 = Document('2학년 광고콘텐츠제작.docx')
    
    # 2학년 정답표 추출
    for table in doc2.tables:
        print(f"\n테이블 발견 (행: {len(table.rows)}, 열: {len(table.columns)})\n")
        
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            print(f"{cells}")
            
except FileNotFoundError:
    print("\n2학년 광고콘텐츠제작.docx 파일을 찾을 수 없습니다.")
